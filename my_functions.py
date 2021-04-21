#!/usr/bin/env python
# coding: utf-8

# In[2]:


# Function to extract channel from ABR log
def abr_channel_received(prefix_hex, file_line, abr_log_received):
    import re
    import logging
    # Regex used to match relevant loglines 
    inmarsat_regex = re.compile("INMARSAT CHANNEL - New packet was received" + r'.*?' + prefix_hex + r'.*?' + file_line)
    gprs_regex = re.compile(" TCP CHANNEL CONNECTION - New packet was received" + r'.*?' + prefix_hex + r'.*?' + file_line)
    temp = 'default'   # Variable to control duplicated messages with same prefix_hex and file_line
    position = -1
    gprs_attempt = 0
    inmarsat_attempt = 0
    
    for line in abr_log_received:
    # If log line matches our regex
        if (gprs_regex.search(line)):
            if temp == 'default':      # When the first line is found, temp is being changed to the last 8 digits of line (see below why)
                gprs_attempt += 1
                logging.info('ABR GPRS: '+str(line))
                temp = line[-9:-1]
                position = abr_log_received.index(line)
            elif temp == line[-9:-1]:  # Once temp is set, it will have the last 8 UNIQUE digits of line, which will identify the difference between messages that have the exact same content and loco id
                gprs_attempt += 1
                logging.info('ABR GPRS: '+str(line))                
        elif (inmarsat_regex.search(line)):
            if temp == 'default':
                inmarsat_attempt += 1
                logging.info('ABR INMARSAT: '+str(line))
                temp = line[-9:-1]
                position = abr_log_received.index(line)   # Position is saved to delete the line from the ABR log (see below why)
            elif temp == line[-9:-1]:
                inmarsat_attempt += 1
                logging.info('ABR INMARSAT: '+str(line))
                
    
    if position != -1:   # Removing the line from the ABR log is necessary so it won't appear again on the next iteration (this can happen with messages that have the same content)
        abr_log_received.pop(position)
        
    
    return(gprs_attempt, inmarsat_attempt)


# In[3]:


# Function to extract channel from ABR log
def abr_channel_sent(prefix_hex, file_line, abr_log_sent):
    import re
    import logging
    # Regex used to match relevant loglines 
    inmarsat_regex = re.compile("INMARSAT CHANNEL - Message submitted to gateway" + r'.*?' + prefix_hex + r'.*?' + file_line)
    gprs_regex = re.compile("TCP CHANNEL CONNECTION - Sending packet" + r'.*?' + prefix_hex + r'.*?' + file_line)
    
    gprs_attempt = 0
    inmarsat_attempt = 0
    
    for line in abr_log_sent:
    # If log line matches our regex
        if (gprs_regex.search(line)):
            gprs_attempt += 1
            logging.info('ABR GPRS: '+str(line))
        elif (inmarsat_regex.search(line)):    
            inmarsat_attempt += 1
            logging.info('ABR INMARSAT: '+str(line))
             

    return(gprs_attempt, inmarsat_attempt)


# In[4]:


# Function for returning the average delivery time of sent, received and total (sent + received) messages
def average_delivery(df_sent,df_received, totalSentMessages, totalReceivedMessages):
    #Packages
    from datetime import timedelta
    from datetime import datetime   
    import numpy as np
    
    # Sent
    df_sent['DELIVERY_TIME'] = df_sent['HM_DAT_OBC']-df_sent['HM_DAT_STC']
    # Converting to string in order to execute a mathematic operation
    df_sent['DELIVERY_TIME'] = df_sent['DELIVERY_TIME']/np.timedelta64(1, 's')
    totalSent = sum(df_sent.iloc[:,12])
    averageSentDelivery = round(totalSent/totalSentMessages,0)
    # Converting from float to mm:ss
    averageSentDelivery = timedelta(seconds=averageSentDelivery)
    averageSentDelivery=str(averageSentDelivery)

    # Received
    df_received['DELIVERY_TIME'] = df_received['HM_DAT_STC']-df_received['HM_DAT_OBC']
    # Converting to string in order to execute a mathematic operation
    df_received['DELIVERY_TIME'] = df_received['DELIVERY_TIME']/np.timedelta64(1, 's')
    totalReceived = sum(df_received.iloc[:,12])
    averageReceivedDelivery = round(totalReceived/totalReceivedMessages,0)
    # Converting from float to mm:ss
    averageReceivedDelivery = timedelta(seconds=averageReceivedDelivery)
    averageReceivedDelivery = str(averageReceivedDelivery)
    
    # Return the average delivery time of Sent and Received messages combined
    averageDelivery = round((totalSent + totalReceived)/(totalSentMessages + totalReceivedMessages),0)
    # Converting from float to mm:ss
    averageDelivery = timedelta(seconds=averageDelivery)
    averageDelivery = str(averageDelivery)


    return(averageSentDelivery, averageReceivedDelivery, averageDelivery)


# In[5]:


# Function to calculate the average delay per SB
def aver_delay_sb(df_sb):
    # Packages
    from datetime import timedelta
    
    # Sent
    for i in range(0,len(df_sb['NAME']),1):
        if df_sb.iloc[i,4] != 0:
            df_sb.iloc[i,10] = df_sb.iloc[i,7]/df_sb.iloc[i,4]
            df_sb.iloc[i,10] = round(df_sb.iloc[i,10],0)    # Have to round the values to Integer to avoid having milliseconds
            df_sb.iloc[i,10] = str(timedelta(seconds=df_sb.iloc[i,10])) # Converting to String time format
            df_sb.iloc[i,7] = str(timedelta(seconds=df_sb.iloc[i,7]))
        else:
            df_sb.iloc[i,10] = 0      
            df_sb.iloc[i,10] = str(timedelta(seconds=df_sb.iloc[i,10]))
            df_sb.iloc[i,7] = str(timedelta(seconds=df_sb.iloc[i,7]))

    # Received
    for i in range(0,len(df_sb['NAME']),1):
        if df_sb.iloc[i,5] != 0:
            df_sb.iloc[i,11] = df_sb.iloc[i,8]/df_sb.iloc[i,5]
            df_sb.iloc[i,11] = round(df_sb.iloc[i,11],0)    # Have to round the values to Integer to avoid having milliseconds
            df_sb.iloc[i,11] = str(timedelta(seconds=df_sb.iloc[i,11])) # Converting to String time format
            df_sb.iloc[i,8] = str(timedelta(seconds=df_sb.iloc[i,8]))
        else:
            df_sb.iloc[i,11] = 0      
            df_sb.iloc[i,11] = str(timedelta(seconds=df_sb.iloc[i,11]))
            df_sb.iloc[i,8] = str(timedelta(seconds=df_sb.iloc[i,8]))

    # Total
    for i in range(0,len(df_sb['NAME']),1):
        if df_sb.iloc[i,6] != 0:
            df_sb.iloc[i,12] = df_sb.iloc[i,9]/df_sb.iloc[i,6]
            df_sb.iloc[i,12] = round(df_sb.iloc[i,12],0)    # Have to round the values to Integer to avoid having milliseconds
            df_sb.iloc[i,12] = str(timedelta(seconds=df_sb.iloc[i,12])) # Converting to String time format
            df_sb.iloc[i,9] = str(timedelta(seconds=df_sb.iloc[i,9]))
        else:
            df_sb.iloc[i,12] = 0      
            df_sb.iloc[i,12] = str(timedelta(seconds=df_sb.iloc[i,12]))
            df_sb.iloc[i,9] = str(timedelta(seconds=df_sb.iloc[i,9]))
            
    return(df_sb)


# In[6]:


# Function to create horizontal barplot for sent, received and total SBs combined
def barplot_all(df_sb, top10_Sent, top10_Received, top10_Total, month, year, images_location):
    # Packages
    import matplotlib.pyplot as plt
    
    plt.figure(1)
    plt.subplot(2,2,1)
    plt.barh(top10_Sent.iloc[:,0],top10_Sent.iloc[:,4], color ='cornflowerblue',align='center')
    plt.xlim(0,max(df_sb.iloc[:,4])+10)
    plt.xlabel("Number of messages") 
    plt.ylabel("Top 10 Section Blocks")
    plt.title("Sent Messages " + month + '-' + year) 
    # Adding percentage values of '%' column to the bars
    #for index, value in enumerate(top10_Sent.iloc[:,4]):
    #    plt.text(value, index, str(top10_Sent.iloc[index,13]))
    plt.subplot(2,2,2)
    plt.barh(top10_Received.iloc[:,0],top10_Received.iloc[:,5], color ='tab:red',align='center')
    plt.xlim(0,max(df_sb.iloc[:,5])+10)
    plt.title("Received Messages " + month + '-' + year) 
    plt.subplot(2,2,3)
    plt.barh(top10_Total.iloc[:,0],top10_Total.iloc[:,6], color ='olivedrab',align='center')
    plt.xlim(0,max(df_sb.iloc[:,6])+10)
    plt.title("All Messages " + month + '-' + year) 
    plt.savefig(images_location+'/bar_all'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[7]:


# Function to display the top 10 locomotives with delay
def barplot_loco(df_loco, month, year, images_location):
    #Packages
    import matplotlib.pyplot as plt
    
    ordered_loco = df_loco.sort_values(by=['%_DELAY'], ascending=False)
    top10_loco = ordered_loco.head(n=10)
    top10_loco = top10_loco.sort_values(by=['%_DELAY'], ascending=True)

    plt.barh(top10_loco.iloc[:,0],top10_loco.iloc[:,6], color ='gold',align='center')
    plt.xlim(0,max(df_loco.iloc[:,6])+200)
    plt.xlabel("Number of messages") 
    plt.ylabel("Locomotives")
    plt.title("Delays per Locomotive " + month + '-' + year) 
    # Adding percentage values of '%' column to the bars
    for index, value in enumerate(top10_loco.iloc[:,6]):
        plt.text(value, index, str(top10_loco.iloc[index,7]))
    plt.savefig(images_location+'/bar_loco'+'_'+ month + '_' + year +'.png')
    plt.show()
    


# In[8]:


# Function to create horizontal barplot for received SBs
def barplot_received(df_sb, top10_Received, month, year, images_location):
    # Packages
    import matplotlib.pyplot as plt
    
    plt.barh(top10_Received.iloc[:,0],top10_Received.iloc[:,5], color ='tab:red',align='center')
    plt.xlim(0,max(df_sb.iloc[:,5])+10)
    plt.xlabel("Number of messages") 
    plt.ylabel("Section Blocks")
    plt.title("Received messages with >5min. delivery time " + month + '-' + year) 
    # Adding percentage values of '%' column to the bars
    for index, value in enumerate(top10_Received.iloc[:,5]):
        plt.text(value, index, str(top10_Received.iloc[index,14]))
    plt.savefig(images_location+'/bar_received'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[9]:


# Function to create horizontal barplot for sent SBs
def barplot_sent(df_sb, top10_Sent, month, year, images_location):
    # Packages
    import matplotlib.pyplot as plt
    
    plt.barh(top10_Sent.iloc[:,0],top10_Sent.iloc[:,4], color ='cornflowerblue',align='center')
    plt.xlim(0,max(df_sb.iloc[:,4])+10)
    plt.xlabel("Number of messages") 
    plt.ylabel("Section Blocks")
    plt.title("Sent messages with >5min. delivery time " + month + '-' + year) 
    # Adding percentage values of '%' column to the bars
    for index, value in enumerate(top10_Sent.iloc[:,4]):
        plt.text(value, index, str(top10_Sent.iloc[index,13]))
    plt.savefig(images_location+'/bar_sent'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[10]:


# Function to create horizontal barplot for total SBs (sent + received)
def barplot_total(df_sb, top10_Total, month, year, images_location):
    # Packages
    import matplotlib.pyplot as plt
    
    plt.barh(top10_Total.iloc[:,0],top10_Total.iloc[:,6], color ='olivedrab',align='center')
    plt.xlim(0,max(df_sb.iloc[:,6])+10)
    plt.xlabel("Number of messages") 
    plt.ylabel("Section Blocks")
    plt.title("Number of Delays per Section Block " + month + '-' + year) 
    # Adding percentage values of '%' column to the bars
    for index, value in enumerate(top10_Total.iloc[:,6]):
        plt.text(value, index, str(top10_Total.iloc[index,15]))
    plt.savefig(images_location+'/bar_total'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[11]:


# Function for returning % and total amount of messages that took more than 5min for both sent, received and total(sent + received) messages.
def calc_delay(df_sent, df_received, totalSentMessages, totalReceivedMessages, a):
    # Sent
    totalSentDelay= df_sent.loc[(df_sent['HM_DAT_OBC']-df_sent['HM_DAT_STC']) >= a]
    totalSentDelay = len(totalSentDelay)
    percentageSentDelay=(totalSentDelay/totalSentMessages)*100
    percentageSentDelay = round(percentageSentDelay,2)
    percentageSentDelay = str(percentageSentDelay)+ "%"

    # Received
    totalReceivedDelay= df_received.loc[(df_received['HM_DAT_STC']-df_received['HM_DAT_OBC']) >= a]
    totalReceivedDelay = len(totalReceivedDelay)
    percentageReceivedDelay=(totalReceivedDelay/totalReceivedMessages)*100
    percentageReceivedDelay = round(percentageReceivedDelay,2)
    percentageReceivedDelay = str(percentageReceivedDelay)+ "%"

    # Total
    totalDelay = totalSentDelay + totalReceivedDelay
    totalMessages = totalSentMessages + totalReceivedMessages
    percentageTotalDelay = (totalDelay/totalMessages)*100
    percentageTotalDelay = round(percentageTotalDelay,2)
    percentageTotalDelay = str(percentageTotalDelay)+"%"


    return(percentageSentDelay, totalSentDelay, percentageReceivedDelay, totalReceivedDelay, percentageTotalDelay, totalDelay)


# In[12]:


# Call function to return the total of delays per channel and queue delays
def channel_calc_received(df_received, totalReceivedDelay):
    import pandas as pd
    
    gprs_delay_received, inmarsat_delay_received, gprs_inmarsat_received = 0, 0, 0  

    gprs_inmarsat_received = df_received.apply(lambda x: (x.GPRS == 1 & x.SAT == 1), axis=1).sum()   # Returns the number of messages with same amounts of attempts for GPRS and SAT channels (intersection)
    gprs_inmarsat_received = int(round(gprs_inmarsat_received/2))       # Divide by two for both channels (GPRS and SAT) so it gives correct amount of total delayed messages (evenly distributed)



    gprs_delay_received = df_received['GPRS'].sum()-gprs_inmarsat_received    # Reduce the intersection from total number per channel
    inmarsat_delay_received = df_received['SAT'].sum()-gprs_inmarsat_received
       
    perc_gprs_delay_received = str(round((gprs_delay_received/totalReceivedDelay)*100,2)) + '%'
    perc_inmarsat_delay_received = str(round((inmarsat_delay_received/totalReceivedDelay)*100,2)) + '%'
    
    
    return(gprs_delay_received, inmarsat_delay_received, perc_gprs_delay_received, perc_inmarsat_delay_received)


# In[13]:


# Call function to return the total of delays per channel and queue delays
def channel_calc_sent(df_sent, totalSentDelay):
    import pandas as pd
    
    gprs_delay_sent, inmarsat_delay_sent, gprs_queue_sent, inmarsat_queue_sent, gprs_non_queue_sent, inmarsat_non_queue_sent, no_attempt_sent, gprs_inmarsat_sent = 0, 0, 0, 0, 0, 0, 0, 0    # no_attempt represent the number of messages without attempts

    gprs_inmarsat_sent = df_sent.apply(lambda x: (x.GPRS == 1 & x.SAT == 1), axis=1).sum()   # Returns the number of messages with same amounts of attempts for GPRS and SAT channels (intersection)
    gprs_inmarsat_sent = int(round(gprs_inmarsat_sent/2))       # Divide by two for both channels (GPRS and SAT) so it gives correct amount of total delayed messages



    gprs_delay_sent = df_sent['GPRS'].sum()-gprs_inmarsat_sent     # Reduce the intersection from total number per channel
    inmarsat_delay_sent = df_sent['SAT'].sum()-gprs_inmarsat_sent
    no_attempt_sent = df_sent['NaN'].sum() # no_attempt represent the number of messages without attempts
    gprs_queue_sent = df_sent.apply(lambda x: (x.QUEUE_DELAY == 1 & x.GPRS == 1), axis=1).sum()   # Not substracting the intersection because it can only have one attempt (one channel)
    inmarsat_queue_sent = df_sent.apply(lambda x: (x.QUEUE_DELAY == 1 & x.SAT == 1), axis=1).sum()
    gprs_non_queue_sent = gprs_delay_sent - gprs_queue_sent
    inmarsat_non_queue_sent = inmarsat_delay_sent - inmarsat_queue_sent
    
    perc_gprs_delay_sent = str(round((gprs_delay_sent/totalSentDelay)*100,2)) + '%'
    perc_inmarsat_delay_sent = str(round((inmarsat_delay_sent/totalSentDelay)*100,2)) + '%'
    perc_no_attempt_sent = str(round((no_attempt_sent/totalSentDelay)*100,2)) + '%'
    perc_gprs_queue_sent = str(round((gprs_queue_sent/totalSentDelay)*100,2)) + '%'
    perc_gprs_non_queue_sent = str(round((gprs_non_queue_sent/totalSentDelay)*100,2)) + '%'
    perc_inmarsat_queue_sent = str(round((inmarsat_queue_sent/totalSentDelay)*100,2)) + '%'
    perc_inmarsat_non_queue_sent = str(round((inmarsat_non_queue_sent/totalSentDelay)*100,2)) + '%'
    perc_total_queue_sent = str(round(((gprs_queue_sent+inmarsat_queue_sent)/totalSentDelay)*100,2)) + '%'
    perc_total_non_queue_sent = str(round(((gprs_non_queue_sent+inmarsat_non_queue_sent)/totalSentDelay)*100,2)) + '%'
    
    return(gprs_delay_sent, inmarsat_delay_sent, gprs_queue_sent, inmarsat_queue_sent, gprs_non_queue_sent, inmarsat_non_queue_sent, no_attempt_sent, perc_gprs_delay_sent, perc_inmarsat_delay_sent, perc_no_attempt_sent, perc_gprs_queue_sent, perc_gprs_non_queue_sent, perc_inmarsat_queue_sent, perc_inmarsat_non_queue_sent, perc_total_queue_sent, perc_total_non_queue_sent)
    


# In[14]:


# Function to count delays per channel
def channel_count_received(i, df_received, gprs_attempt, inmarsat_attempt):
    
    # Counts the amount of attempts to determine if the delay was caused by GPRS or INMARSAT channel
    if gprs_attempt > inmarsat_attempt:
        df_received.iloc[i,14] += 1
    elif inmarsat_attempt > gprs_attempt: 
        df_received.iloc[i,15] += 1
    elif inmarsat_attempt == gprs_attempt and (inmarsat_attempt+gprs_attempt) != 0: # Equal number of attempts between GPRS and SAT will mark both channels as delay
        df_received.iloc[i,14] += 1  
        df_received.iloc[i,15] += 1 
    
    return(df_received)


# In[15]:


# Function to count delays per channel
def channel_count_sent(i, df_sent, gprs_attempt, inmarsat_attempt):
    
    # Counts the amount of attempts to determine if the delay was caused by GPRS or INMARSAT channel
    if gprs_attempt > inmarsat_attempt:
        df_sent.iloc[i,14] += 1
    elif inmarsat_attempt > gprs_attempt: 
        df_sent.iloc[i,15] += 1
    elif inmarsat_attempt == gprs_attempt and (inmarsat_attempt+gprs_attempt) != 0: # Equal number of attempts between GPRS and SAT will mark both channels as delay
        df_sent.iloc[i,14] += 1  
        df_sent.iloc[i,15] += 1 

    if inmarsat_attempt == 0 and gprs_attempt == 0:
        df_sent.iloc[i,17] += 1  
    elif (inmarsat_attempt == 1 and gprs_attempt == 0) or (inmarsat_attempt == 0 and gprs_attempt == 1):
        df_sent.iloc[i,16] += 1

    
    return(df_sent)


# In[16]:


# Function to retrieve config information
def config_read(month, year):
    import configparser
    import logging
    import os
    import sys
    import cx_Oracle

    
    parser = configparser.ConfigParser()
    parser.read("config.ini")   # Reading config file
    
    db_username = parser.get("DATABASE_IMPORT", "USERNAME")
    db_password = parser.get("DATABASE_IMPORT", "PASSWORD")
    db_host = parser.get("DATABASE_IMPORT", "HOST")
    db_alias = parser.get("DATABASE_IMPORT", "NETWORK_ALIAS")
    dbexp_username = parser.get("DATABASE_EXPORT", "USERNAME")
    dbexp_password = parser.get("DATABASE_EXPORT", "PASSWORD")
    dbexp_host = parser.get("DATABASE_EXPORT", "HOST")
    dbexp_alias = parser.get("DATABASE_EXPORT", "NETWORK_ALIAS")
    abr_log_location = parser.get("LOGS", "ABR_PATH")
    mm_log_location = parser.get("LOGS", "MESSAGE_MANAGER_PATH")
    output_log_location = parser.get("LOGS", "OUTPUT_LOGS")
    images_location = parser.get("IMAGES", "PATH")
    report_location = parser.get("REPORT", "PATH")
    
    logging.basicConfig(level=logging.INFO, filename=output_log_location+"\data_app_"+ month +"_"+ year + ".log", format="%(asctime)-15s %(levelname)-8s %(message)s",force=True)
    logging.info('#### SCRIPT EXECUTION ####')
    logging.info('------------------IMPORT CONFIGFILE------------------------------') 
    logging.info('USERNAME: '+ db_username)
    logging.info('PASSWORD: '+ db_password)
    logging.info('HOST: '+ db_host)
    logging.info('NETWORK ALIAS: '+ db_alias)
    logging.info('USERNAME: '+ dbexp_username)
    logging.info('PASSWORD: '+ dbexp_password)
    logging.info('HOST: '+ dbexp_host)
    logging.info('NETWORK ALIAS: '+ dbexp_alias)    
    logging.info('MESSAGE MANAGER LOG LOCATION: '+ mm_log_location)
    logging.info('ABR LOG LOCATION: '+ abr_log_location)
    logging.info('OUTPUT LOG LOCATION: '+ output_log_location)
    logging.info('IMAGES SAVE LOCATION: '+ images_location)
    logging.info('REPORT SAVE LOCATION: '+ report_location)
    logging.info('---------------END OF IMPORT CONFIGFILE---------------------------')
    
    if not(os.path.exists(mm_log_location)):                      # To make sure MessageManager log exists
        logging.error('!! MESSAGE MANAGER LOG NOT FOUND !!')
        logging.error('Please check the MESSAGE_MANAGER_PATH parameter on the config.ini file and try again.')
        logging.critical('#### SCRIPT TERMINATED ####')
        sys.exit() 
    
    if not(os.path.exists(abr_log_location)):                          # To make sure ABR log exists
        logging.error('!! ABR LOG NOT FOUND !!')
        logging.error('Please check the ABR_PATH parameter on the config.ini file and try again.')
        logging.critical('#### SCRIPT TERMINATED ####')
        sys.exit()     

    if not(os.path.exists(output_log_location)):                          # To make sure output log path exists
        logging.error('!! PATH OUTPUT LOG NOT FOUND !!')
        logging.error('Please check the OUTPUT_LOGS parameter on the config.ini file and try again.')
        logging.critical('#### SCRIPT TERMINATED ####')
        sys.exit()     
        
    if not(os.path.exists(images_location)):                          # To make sure images path exists
        logging.error('!! PATH IMAGES NOT FOUND !!')
        logging.error('Please check the PATH parameter on the config.ini file and try again.')
        logging.critical('#### SCRIPT TERMINATED ####')
        sys.exit() 
        
    if not(os.path.exists(report_location)):                          # To make sure images path exists
        logging.error('!! PATH REPORT NOT FOUND !!')
        logging.error('Please check the PATH parameter on the config.ini file and try again.')
        logging.critical('#### SCRIPT TERMINATED ####')
        sys.exit()
     
    try:
        con = cx_Oracle.connect(db_username+'/'+db_password+'@'+db_host+'/'+db_alias)
        con.close()
    except:
        logging.error('!! IMPORT DATABASE INFO INVALID !!')
        logging.error('Please check the DATABASE_IMPORT tag on the config.ini file and try again.')
        logging.critical('#### SCRIPT TERMINATED ####')
        sys.exit()
    logging.info('Connection with '+db_username+' database successful!')
    try:
        con = cx_Oracle.connect(dbexp_username+'/'+dbexp_password+'@'+dbexp_host+'/'+dbexp_alias)
        con.close()
    except:
        logging.error('!! EXPORT DATABASE INFO INVALID !!')
        logging.error('Please check the DATABASE_IMPORT tag on the config.ini file and try again.')
        logging.critical('#### SCRIPT TERMINATED ####')
        sys.exit()        
    logging.info('Connection with '+dbexp_username+' database successful!')
    
    return(db_username, db_password, db_host, db_alias, dbexp_username, dbexp_password, dbexp_host, dbexp_alias, abr_log_location, mm_log_location, output_log_location, images_location, report_location)


# In[17]:


# Function for returning the total amount of sent messages, received messages and total (sent + received) messages
def count_messages(df_sent, df_received):
    # Sent
    totalSentMessages = len(df_sent['HM_ID_HM'])
    # Received
    totalReceivedMessages = len(df_received['HM_ID_HM'])
    # Total
    totalMessages = totalSentMessages + totalReceivedMessages

    return(totalSentMessages, totalReceivedMessages, totalMessages)


# In[18]:


# Function for database import
def db_import(db_username, db_password, db_host, db_alias):
    #Packages
    import cx_Oracle
    import pandas as pd
    
    # Initializing database connection
    con = cx_Oracle.connect(db_username+'/'+db_password+'@'+db_host+'/'+db_alias)
    cur = con.cursor()

    # First Query
    query=cur.execute('select * from historico_mensagens')
    df = pd.DataFrame(query)
    # Second Query
    query=cur.execute('select EV_NOM_MAC from ELEM_VIA where TE_ID_TP = 3')
    df_sb = pd.DataFrame(query)

    cur.close()
    con.close()
    
    return (df,df_sb)


# In[19]:


# Function for database insertion
def db_insert(month, year, density, dispersion_all, dispersion_sent_received, dispersion_total, bar_all, bar_sent, bar_received, bar_total, bar_loco, time_series, time_series_decompose, totalSentMessages, averageSentDelivery, percentageSentDelay, totalSentDelay, averageSentDeliveryDelay, totalReceivedMessages, averageReceivedDelivery, percentageReceivedDelay, totalReceivedDelay, averageReceivedDeliveryDelay,totalMessages, averageTotalDelivery, percentageTotalDelay, totalDelay, averageDeliveryDelay, df_sb, df_loco, ts_total, gprs_delay_sent, inmarsat_delay_sent, gprs_queue_sent, inmarsat_queue_sent, gprs_non_queue_sent, inmarsat_non_queue_sent, no_attempt_sent, perc_gprs_delay_sent, perc_inmarsat_delay_sent, perc_no_attempt_sent, perc_gprs_queue_sent, perc_gprs_non_queue_sent, perc_inmarsat_queue_sent, perc_inmarsat_non_queue_sent, perc_total_queue_sent, perc_total_non_queue_sent, gprs_delay_received, inmarsat_delay_received, perc_gprs_delay_received, perc_inmarsat_delay_received, dbexp_username, dbexp_password, dbexp_host, dbexp_alias):
    # Packages
    import cx_Oracle
    import pandas as pd
    
    con = cx_Oracle.connect(dbexp_username+'/'+dbexp_password+'@'+dbexp_host+'/'+dbexp_alias)
    cur = con.cursor()

    cur.execute("insert into SENT (ID_SENT, MONTH, YEAR, TOTALSENTMESSAGES, AVERAGESENTDELIVERY, PERCENTAGESENTDELAY, TOTALSENTDELAY, averageSentDeliveryDelay, gprs_delay_sent, perc_gprs_delay_sent, gprs_non_queue_sent, perc_gprs_non_queue_sent, gprs_queue_sent, perc_gprs_queue_sent, inmarsat_delay_sent, perc_inmarsat_delay_sent, inmarsat_non_queue_sent, perc_inmarsat_non_queue_sent, inmarsat_queue_sent, perc_inmarsat_queue_sent, no_attempt_sent, perc_no_attempt_sent, perc_total_queue_sent, perc_total_non_queue_sent, regtimestamp) values (SENT_ID_SEQ.NEXTVAL, :month, :year, :totalSentMessages, :averageSentDelivery, :percentageSentDelay, :totalSentDelay, :averageSentDeliveryDelay, :gprs_delay_sent, :perc_gprs_delay_sent, :gprs_non_queue_sent, :perc_gprs_non_queue_sent, :gprs_queue_sent, :perc_gprs_queue_sent, :inmarsat_delay_sent, :perc_inmarsat_delay_sent, :inmarsat_non_queue_sent, :perc_inmarsat_non_queue_sent, :inmarsat_queue_sent, :perc_inmarsat_queue_sent, :no_attempt_sent, :perc_no_attempt_sent, :perc_total_queue_sent, :perc_total_non_queue_sent, sysdate)", [month, year, totalSentMessages, averageSentDelivery, percentageSentDelay, totalSentDelay, averageSentDeliveryDelay, int(gprs_delay_sent), perc_gprs_delay_sent, int(gprs_non_queue_sent), perc_gprs_non_queue_sent, int(gprs_queue_sent), perc_gprs_queue_sent, int(inmarsat_delay_sent), perc_inmarsat_delay_sent, int(inmarsat_non_queue_sent), perc_inmarsat_non_queue_sent, int(inmarsat_queue_sent), perc_inmarsat_queue_sent, int(no_attempt_sent), perc_no_attempt_sent, perc_total_queue_sent, perc_total_non_queue_sent])
    cur.execute("insert into RECEIVED (ID_RECEIVED, MONTH, YEAR, TOTALRECEIVEDMESSAGES, AVERAGERECEIVEDDELIVERY, PERCENTAGERECEIVEDDELAY, TOTALRECEIVEDDELAY, averageReceivedDeliveryDelay, gprs_delay_received, perc_gprs_delay_received, inmarsat_delay_received, perc_inmarsat_delay_received, regtimestamp) values (RECEIVED_ID_SEQ.NEXTVAL, :month, :year, :totalReceivedMessages, :averageReceivedDelivery, :percentageReceivedDelay, :totalReceivedDelay, :averageReceivedDeliveryDelay, :gprs_delay_received, :perc_gprs_delay_received, :inmarsat_delay_received, :perc_inmarsat_delay_received, sysdate)", [month, year, totalReceivedMessages, averageReceivedDelivery, percentageReceivedDelay, totalReceivedDelay, averageReceivedDeliveryDelay, int(gprs_delay_received), perc_gprs_delay_received, int(inmarsat_delay_received), perc_inmarsat_delay_received])
    cur.execute("insert into TOTAL (ID_TOTAL, MONTH, YEAR, TOTALMESSAGES, AVERAGETOTALDELIVERY, PERCENTAGETOTALDELAY, TOTALDELAY, averageDeliveryDelay, regtimestamp) values (TOTAL_ID_SEQ.NEXTVAL, :month, :year, :totalMessages, :averageTotalDelivery, :percentageTotalDelay, :totalDelay, :averageDeliveryDelay, sysdate)", [month, year, totalMessages, averageTotalDelivery, percentageTotalDelay, totalDelay, averageDeliveryDelay])
    cur.execute("insert into GRAPHS (ID_GRAPHS, MONTH, YEAR, regtimestamp) values (GRAPHS_ID_SEQ.NEXTVAL, :month, :year, sysdate)", [month, year])
    cur.callproc("graphs_pr",[month, year, density, dispersion_all, dispersion_sent_received, dispersion_total, bar_all, bar_sent, bar_received, bar_total, bar_loco, time_series, time_series_decompose])
    for i in range(0,len(df_sb['NAME']),1):
        cur.execute("insert into DF_SB (ID_DF_SB, NAME, MONTH, YEAR, counts_sent, counts_received, counts_total, average_time_sent, average_time_received, average_time, perc_sent, perc_received, perc_total, regtimestamp) values (DF_SB_ID_SEQ.NEXTVAL, :name , :month, :year, :counts_sent, :counts_received, :counts_total,:average_time_sent, :average_time_received, :average_time, :perc_sent, :perc_received, :perc_total, sysdate)", [df_sb.iloc[i,0],month, year, int(df_sb.iloc[i,4]), int(df_sb.iloc[i,5]), int(df_sb.iloc[i,6]), df_sb.iloc[i,10], df_sb.iloc[i,11], df_sb.iloc[i,12], df_sb.iloc[i,13], df_sb.iloc[i,14], df_sb.iloc[i,15]])   
    for i in range(0,len(df_loco['HM_LOCO_MSG']),1):
        cur.execute("insert into LOCO (ID_LOCO, HM_LOCO_MSG, MONTH, YEAR, T_SENT, T_RECEIVED, TOTAL, D_SENT, D_RECEIVED, D_TOTAL, PERC_DELAY, regtimestamp) values (LOCO_ID_SEQ.NEXTVAL, :HM_LOCO_MSG , :month, :year, :T_SENT, :T_RECEIVED, :TOTAL, :D_SENT, :D_RECEIVED, :D_TOTAL, :PERC_DELAY, sysdate)", [df_loco.iloc[i,0],month, year, int(df_loco.iloc[i,1]), int(df_loco.iloc[i,2]), int(df_loco.iloc[i,3]), int(df_loco.iloc[i,4]), int(df_loco.iloc[i,5]), int(df_loco.iloc[i,6]), df_loco.iloc[i,7]])   
    for i in range(0,len(ts_total['TS_DATE']),1):  
        cur.execute("insert into TIME_SERIES (ID_TS, MONTH, YEAR, TS_DATE, SENT, RECEIVED, TOTAL, regtimestamp) values (TS_ID_SEQ.NEXTVAL, :month, :year, :ts_date, :sent, :received, :total, sysdate)", [month, year, ts_total.iloc[i,0], int(ts_total.iloc[i,1]), int(ts_total.iloc[i,2]), int(ts_total.iloc[i,3])])
    
    con.commit() 
    cur.close()
    con.close() 


# In[20]:


# Function to create Density graph
def dens_graph(df_sent, df_received, month, year, images_location):
    #Packages
    import seaborn as sns
    import matplotlib.pyplot as plt
    
    sns.distplot(df_sent['DELIVERY_TIME'],hist=False,kde=True,color = 'cornflowerblue', label = 'Sent')
    sns.distplot(df_received['DELIVERY_TIME'],hist=False,kde=True, color = 'tab:red', label = 'Received')
    plt.title('Delivery times '+month+'-'+year)
    plt.xlabel("Delivery Times (in seconds)") 
    plt.legend(loc='best')
    plt.xlim(0, 500)
    plt.savefig(images_location+'/density'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[21]:


# Function for formatting dataframe
def df_format(df, df_sb):

    # Adding names to the columns of the dataframes
    df.columns=['HM_ID_HM','HM_LOCO_MSG','HM_PRF_TRM','HM_OBC_MSG','HM_NUM_MSG','HM_TP_MSG','HM_DAT_OBC','HM_DAT_STC',
                'HM_TXT_MSG','HM_TXT_TW','HM_LAT_HM','HM_LON_HM','HM_ENV_RCB','HM_MSG_VOZ','HM_NUM_FUN','HM_ID_AUX','HM_SIT_ENV','MT_ID_MT']
    #df_sb.columns=['NAME']
    # Dropping tables with no use
    df = df.drop(['HM_TP_MSG','HM_TXT_TW','HM_ID_AUX','HM_OBC_MSG','HM_LAT_HM','HM_LON_HM'], axis=1)
    df_sb.columns=['NAME'] 
    for i in range(0,len(df_sb['NAME']),1):
        df_sb['SECTION_DEBUT'] = ("Section Début: " + df_sb.loc[:,'NAME'] + " ")
        df_sb['CANTON'] = "Canton: " + df_sb.loc[:,'NAME'] + " "
        df_sb['SECTION_TETE'] = ("Section da Tête: " + df_sb.loc[:,'NAME'] + " ")
    df_sb['COUNTS_SENT'] = 0
    df_sb['COUNTS_RECEIVED'] = 0
    df_sb['COUNTS_TOTAL'] = 0
    df_sb['TOTAL_TIME_SENT'] = 0.0
    df_sb['TOTAL_TIME_RECEIVED'] = 0.0
    df_sb['TOTAL_TIME'] = 0.0
    df_sb['AVERAGE_TIME_SENT'] = 0.0
    df_sb['AVERAGE_TIME_RECEIVED'] = 0.0
    df_sb['AVERAGE_TIME'] = 0.0
    df_sb['%_SENT'] = 0
    df_sb['%_RECEIVED'] = 0
    df_sb['%_TOTAL'] = 0
    
    return(df, df_sb)


# In[22]:


# Function to prepare dataframe for channel informations
def df_prep_channel(df_sent, df_received):

    df_sent['GPRS'] = int(0)
    df_sent['SAT'] = int(0)
    df_sent['QUEUE_DELAY'] = 0
    df_sent['NaN'] = 0
    df_sent['MT_ID_MT'] = df_sent['MT_ID_MT'].astype('int')
    
    df_received['GPRS'] = int(0)
    df_received['SAT'] = int(0)
    
    
    return(df_sent, df_received)


# In[23]:


# Function for splitting dataframe in 'Sent' and'Received' messages and locomotives
def df_split(df):
    #Packages
    import pandas as pd
    
    df_sent = df.loc[df['HM_ENV_RCB']=='E']
    df_received = df.loc[df['HM_ENV_RCB']=='R']
    df_loco = pd.DataFrame(columns = ['HM_LOCO_MSG','T_SENT', 'T_RECEIVED','TOTAL','D_SENT','D_RECEIVED','D_TOTAL','%_DELAY'])  # Creating Loco dataframe
    df_loco.iloc[:,0] = df.HM_LOCO_MSG.unique() # Retrieve unique locomotive values
    df_loco['D_SENT'] = 0      # Initializing columns for counting
    df_loco['D_RECEIVED'] = 0
    
    # Removing the original dataframe
    df = pd.DataFrame(None)
    
    return(df, df_sent, df_received, df_loco)


# In[24]:


# Function to create Density graph 
def disp_all_graph(df_sb, month, year, images_location):
    #Packages
    import matplotlib.pyplot as plt
    
    plt.figure(1)
    plt.subplot(2,2,1)
    plt.scatter(df_sb.iloc[:,0],df_sb.iloc[:,4], color='cornflowerblue', marker='.')
    plt.xticks([])
    for i in range(0,max(df_sb.iloc[:,4]),5):
        plt.axhline(y=i, color='black', linestyle='--',linewidth=0.5)
    plt.title("Sent Messages " + month + '-' + year)
    plt.subplot(2,2,2)
    plt.scatter(df_sb.iloc[:,0],df_sb.iloc[:,5], color='tab:red', marker='.')
    plt.xticks([])
    for i in range(0,max(df_sb.iloc[:,5]),10):
        plt.axhline(y=i, color='black', linestyle='--',linewidth=0.5)
    plt.title("Received Messages " + month + '-' + year)
    plt.subplot(2,2,3)
    plt.scatter(df_sb.iloc[:,0],df_sb.iloc[:,6], color='olivedrab', marker='.')
    plt.xticks([])
    for i in range(0,max(df_sb.iloc[:,6]),10):
        plt.axhline(y=i, color='black', linestyle='--',linewidth=0.5)
    plt.title("All Messages " + month + '-' + year) 
    plt.savefig(images_location+'/dispersion_all'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[25]:


# Function to create Density graph of sent and received messages combined
def disp_sent_received_graph(df_sb, month, year, images_location):
    #Packages
    import matplotlib.pyplot as plt
    
    plt.scatter(df_sb.iloc[:,0],df_sb.iloc[:,4], color='cornflowerblue', marker='.', label='Sent')
    plt.xticks([])
    for i in range(0,max(df_sb.iloc[:,4]),10):
        plt.axhline(y=i, color='black', linestyle='--',linewidth=0.5)

    plt.scatter(df_sb.iloc[:,0],df_sb.iloc[:,5], color='tab:red', marker='.', label='Received')
    plt.xticks([])
    for i in range(0,max(df_sb.iloc[:,5]),10):
        plt.axhline(y=i, color='black', linestyle='--',linewidth=0.5)
    plt.legend(loc='best')
    plt.title("Delivery of >5 min. per Section Block " + month + '-' + year)
    plt.savefig(images_location+'/dispersion_sent_received'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[26]:


# Function to create Density graph of total messages
def disp_total_graph(df_sb, month, year, images_location):
    #Packages
    import matplotlib.pyplot as plt
    
    plt.scatter(df_sb.iloc[:,0],df_sb.iloc[:,6], color='olivedrab', marker='.')
    plt.xticks([])
    for i in range(0,max(df_sb.iloc[:,6]),10):
        plt.axhline(y=i, color='black', linestyle='--',linewidth=0.5)
    plt.title("Delivery of >5 min. per Section Block " + month + '-' + year)
    plt.savefig(images_location+'/dispersion_total'+'_'+ month + '_' + year +'.png')
    plt.show()


# In[27]:


# Function for filtering (part I)
def filter_I (df_sent,df_received,month,year):
    # Sent DataFrame
    # 1) Filter messages within the date defined by user
    df_sent = df_sent.loc[df_sent['HM_DAT_STC']>= year + '-' + month +'-01']
    if(month == '12'):     # December
        df_sent = df_sent.loc[df_sent['HM_DAT_STC']< str(int(year)+1) + '-1-01']
    else:                  # Other months
        df_sent = df_sent.loc[df_sent['HM_DAT_STC']< year + '-' + str(int(month)+1) +'-01']
    # 2) Filter all messages that contain OBC time
    df_sent = df_sent.loc[df_sent['HM_DAT_OBC'].notnull()]
    # 3) Filter messages to ensure that the STC time is smaller (or igual) to the OBC time (avoid negative values)
    df_sent = df_sent.loc[df_sent['HM_DAT_STC']<=df_sent['HM_DAT_OBC']]
    # 4) Filter messages that contain locomotive ID
    df_sent = df_sent[df_sent['HM_LOCO_MSG'].notnull()]

    # Received DataFrame
    # 1) Filter all messages that did not expired
    df_received = df_received.loc[df_received['HM_SIT_ENV']=='E']
    # 2) Filter messages from within the date set by user
    df_received = df_received.loc[df_received['HM_DAT_STC']>= year + '-' + month +'-01']
    if(month == '12'):     # December
        df_received = df_received.loc[df_received['HM_DAT_STC']< str(int(year)+1) + '-1-01']
    else:                  # Other months
        df_received = df_received.loc[df_received['HM_DAT_STC']< year + '-' + str(int(month)+1) +'-01']
    # 3) Filter all messages that contain OBC time
    df_received = df_received.loc[df_received['HM_DAT_OBC'].notnull()]
    # 4) Filter all messages that have a date later than needed (avoiding future date messages)
    if(month == '12'):
        df_received = df_received.loc[df_received['HM_DAT_OBC']< str(int(year)+1) +'-1-02']
    else:
        df_received = df_received.loc[df_received['HM_DAT_OBC']< year + '-' + str(int(month)+1) +'-02']
    # 5) Filter messages to ensure that the OBC time is smaller (or igual) to the STC time (avoid negative values)
    df_received = df_received.loc[df_received['HM_DAT_OBC']<=df_received['HM_DAT_STC']]
    # 6) Filter messages that contain locomotive ID
    df_received = df_received[df_received['HM_LOCO_MSG'].notnull()]
  
    
    return(df_sent, df_received)


# In[28]:


# Function for filtering (part II)
def filter_II(df_sent, df_received, a):
    # Sent
    # 3) Filter messages that took longer than 5 minutes
    df_sent = df_sent.loc[(df_sent['HM_DAT_OBC']-df_sent['HM_DAT_STC']) >= a]
    # 4) Order the dataframe by HM_ID_HM
    df_sent = df_sent.sort_values(by=['HM_ID_HM'])

    #Received
    # 5) Filter messages that took longer than 5 minutes
    df_received = df_received.loc[(df_received['HM_DAT_STC']-df_received['HM_DAT_OBC']) >= a]
    # 6) Order the dataframe by HM_ID_HM
    df_received = df_received.sort_values(by=['HM_ID_HM'])
    
    return(df_sent, df_received)


# In[29]:


# Function for filtering (Part III)
def filter_III(df_sent, df_received):
    # Sent
    # 5) Filter messages with macros 1001, 1002 and 1003
    df_sent = df_sent.loc[(df_sent['HM_NUM_MSG']==1001)|(df_sent['HM_NUM_MSG']==1002)|(df_sent['HM_NUM_MSG']==1003)]
    
    # Received
    # 7) Filter messages with macros 2006, 2011, 2012, 2015, 2017, 2020, 2021 and 2029
    df_received = df_received.loc[(df_received['HM_NUM_MSG']==2006)|(df_received['HM_NUM_MSG']==2011)|(df_received['HM_NUM_MSG']==2012)|(df_received['HM_NUM_MSG']==2015)|(df_received['HM_NUM_MSG']==2017)|(df_received['HM_NUM_MSG']==2020)|(df_received['HM_NUM_MSG']==2021)|(df_received['HM_NUM_MSG']==2029)]
    # 8) Remove messages that don't contain data
    df_received = df_received.loc[df_received['HM_TXT_MSG'].notnull()]
    
    return(df_sent, df_received)    


# In[30]:


def import_log(abr_log_location, mm_log_location):
    # Message Manager Import
    myfile = open(mm_log_location, errors="ignore")
    mm_log = myfile.read().splitlines()
    myfile.close()
    
    # ABR import
    abr_log_sent=[]
    abr_log_received=[]
    with open(abr_log_location, "r") as abr:
        for line in abr:
            if (('TCP CHANNEL CONNECTION - Sending packet' in line) or ('INMARSAT CHANNEL - Message submitted to gateway' in line)) and ('11-20-20-30' in line):
                abr_log_sent.append(line)   # adding an element to the end of the list
            elif (('TCP CHANNEL CONNECTION - New packet was received' in line) or ('INMARSAT CHANNEL - New packet was received' in line)) and ('11-20-20-30' in line):
                abr_log_received.append(line)   # adding an element to the end of the list                

    #with open('D:/Logs/2021_2_ABR.log', "r") as abr:
        #for line_2 in abr:
            #elif (('TCP CHANNEL CONNECTION - New packet was received' in line_2) or ('INMARSAT CHANNEL - New packet was received' in line_2)) and ('11-20-20-30' in line_2):
                #abr_log_received.append(line)   # adding an element to the end of the list
    
    return(mm_log, abr_log_sent, abr_log_received)


# In[31]:


# Function to get the number of sent and received messages for each locomotive
def loco_count(df_loco, df_sent, df_received):

    for i in range(0, len(df_loco['HM_LOCO_MSG']),1):
        df_loco.iloc[i,1] = (df_sent['HM_LOCO_MSG'] == df_loco.iloc[i,0]).sum()
        df_loco.iloc[i,2] = (df_received['HM_LOCO_MSG'] == df_loco.iloc[i,0]).sum()  
        df_loco.iloc[i,3] = df_loco.iloc[i,1] + df_loco.iloc[i,2]

    df_loco = df_loco[df_loco['TOTAL']!=0]  # Dropping rows without any messages
        
    
    return(df_loco)


# In[32]:


# Function to return the number of sent, received and total messages with delay
def loco_delay_count(df_loco, df_sent,df_received, a):
    #Packages       
    from datetime import timedelta    
    a = timedelta.total_seconds(a)    # Converting to float to use it for comparison.
    
    for i in range(0, len(df_loco['HM_LOCO_MSG']),1):
        for j in range(0,len(df_sent['HM_ID_HM']),1):
            if (df_sent.iloc[j,1] == df_loco.iloc[i,0]) and (df_sent.iloc[j,12]>=a):
                df_loco.iloc[i,4] += 1
                
    for i in range(0, len(df_loco['HM_LOCO_MSG']),1):
        for j in range(0,len(df_received['HM_ID_HM']),1):
            if (df_received.iloc[j,1] == df_loco.iloc[i,0]) and (df_received.iloc[j,12]>=a):
                df_loco.iloc[i,5] += 1   
                
    df_loco.iloc[:,6] = df_loco.iloc[:,4] + df_loco.iloc[:,5]            
    
    return(df_loco)


# In[33]:


# Function to remove anything from the line that isn't hexadecimal
def mm_hex(file_line):
    # Get the Hex from the MM message
    file_line = file_line[file_line.find('Hex: ')+5:]
    # Removing any duplicated 'FF'
    file_line = file_line.replace("-FF","-")
    # Removing any spaces
    file_line = file_line.strip()
    
    return(file_line)


# In[34]:


# Function that searches for loglines on the MM log for RECEIVED messages
def mm_search_received(i, df_received, new_format, mm_log, new_format_I, new_format_II):
    import re
    import logging
    # Regex used to match relevant loglines 
    line_regex = re.compile(new_format + r'.*?' + " Mensagem Recebida -  Locomotiva: "+df_received.iloc[i,1]+" Tipo de comunica" + r'.*?' + " Hex: *")
    
    file_line = 0 # Set variable as zero to control flow

    for line in mm_log:
        if (line_regex.search(line)):
            file_line = line
            break
         
         # Alternative flow  -  In case line is not found, look at previous minute (08:30 -> 08:29)
    if file_line == 0:
        line_regex_alternative_I = re.compile(new_format_I + r'.*?' + " Mensagem Recebida -  Locomotiva: "+df_received.iloc[i,1]+" Tipo de comunica" + r'.*?' + " Hex: *")
        for line in mm_log:        
            if (line_regex_alternative_I.search(line)):
                file_line = line
                logging.info('Alternative Flow I: '+str(line_regex_alternative_I))
                break

    if file_line == 0:
        # Alternative flow  -  In case line is not found, look at previous hour (08:00 -> 07:59)
        line_regex_alternative_II = re.compile(new_format_II + r'.*?' + " Mensagem Recebida -  Locomotiva: "+df_received.iloc[i,1]+" Tipo de comunica" + r'.*?' + " Hex: *")
        for line in mm_log:         
            if (line_regex_alternative_II.search(line)):
                file_line = line
                logging.info('Alternative Flow II: '+str(line_regex_alternative_II))
                break
       
    return(file_line)


# In[35]:


# Function that searches for loglines on the MM log for SENT messages
def mm_search_sent(i, df_sent, new_format, mm_log, new_format_I, new_format_II):
    import re
    import logging
    # Regex used to match relevant loglines 
    line_regex = re.compile(new_format + r'.*?' + " Mensagem Enviada - Locomotiva: "+df_sent.iloc[i,1]+" Tipo de comunica" + r'.*?' + "mero Mensagem: " + str(df_sent.iloc[i,3])+" Hex: *")
    
    file_line = 0 # Set variable as zero to control flow

    for line in mm_log:
        if (line_regex.search(line)):
            file_line = line
            break
         
         # Alternative flow  -  In case line is not found, look at previous minute (08:30 -> 08:29)
    if file_line == 0:
        line_regex_alternative_I = re.compile(new_format_I + r'.*?' + " Mensagem Enviada - Locomotiva: "+df_sent.iloc[i,1]+" Tipo de comunica" + r'.*?' + "mero Mensagem: " +str(df_sent.iloc[i,3])+" Hex: *")
        for line in mm_log:        
            if (line_regex_alternative_I.search(line)):
                file_line = line
                logging.info('Alternative Flow I: '+str(line_regex_alternative_I))
                break

    if file_line == 0:
        # Alternative flow  -  In case line is not found, look at previous hour (08:00 -> 07:59)
        line_regex_alternative_II = re.compile(new_format_II + r'.*?' + " Mensagem Enviada - Locomotiva: "+df_sent.iloc[i,1]+" Tipo de comunica" + r'.*?' + "mero Mensagem: " +str(df_sent.iloc[i,3])+" Hex: *")
        for line in mm_log:         
            if (line_regex_alternative_II.search(line)):
                file_line = line
                logging.info('Alternative Flow II: '+str(line_regex_alternative_II))
                break
       
    return(file_line)


# In[36]:


# Function to create 3 different time formats for MM log analysis
def mm_timeformat(i, df_sent):    
    # Converting string to MM date format (mm/dd/yyyy)
    new_format=str(df_sent.iloc[i,5])
    temp=""

    for i in new_format: 
        if i.isalnum(): # Strips all special characters
            temp += i

    new_format=temp[4:6]+'/'+temp[6:8]+'/'+temp[0:4]+' '+temp[8:10]+':'+temp[10:12]+':' # Seconds is not included because there may be a difference between DB and MM

    # To cover in case of difference in minutes (08:30 -> 08:29)
    if (temp[10] != '0') and (temp[10:12] != '10'):
        new_format_I = temp[4:6]+'/'+temp[6:8]+'/'+temp[0:4]+' '+temp[8:10]+':'+str(int(temp[10:12])-1)+':'
    else:
        new_format_I = temp[4:6]+'/'+temp[6:8]+'/'+temp[0:4]+' '+temp[8:10]+':0'+str(int(temp[10:12])-1)+':' # Without this it will change an hour '08' into '7', which doesn't work

    # To cover in case of difference in hour (08:00 -> 07:59)
    if temp[8] != '0':
        new_format_II = temp[4:6]+'/'+temp[6:8]+'/'+temp[0:4]+' '+str(int(temp[8:10])-1)+':'+'59'+':' 
    elif temp[8] == '0' and temp[9] != '0':
        new_format_II = temp[4:6]+'/'+temp[6:8]+'/'+temp[0:4]+' 0'+str(int(temp[8:10])-1)+':'+'59'+':' # Without this it will change an hour '08' into '7', which doesn't work
    else:
        new_format_II = temp[4:6]+'/'+str(int(temp[6:8])-1)+'/'+temp[0:4]+' 23:59:' # In case the hour is 00:00:00


    return (new_format, new_format_I, new_format_II)


# In[37]:


# Function to select TOP 10 for barplot graph
def order_barplot(df_sb):
    # Ordering values
    ordered_Sent = df_sb.sort_values(by=['COUNTS_SENT'], ascending=False)
    ordered_Received = df_sb.sort_values(by=['COUNTS_RECEIVED'], ascending=False)
    ordered_Total = df_sb.sort_values(by=['COUNTS_TOTAL'], ascending=False)

    top10_Sent = ordered_Sent.head(n=10)
    top10_Received = ordered_Received.head(n=10)
    top10_Total = ordered_Total.head(n=10)

    top10_Sent = top10_Sent.sort_values(by=['COUNTS_SENT'], ascending=True)
    top10_Received = top10_Received.sort_values(by=['COUNTS_RECEIVED'], ascending=True)
    top10_Total = top10_Total.sort_values(by=['COUNTS_TOTAL'], ascending=True)
    
    return(top10_Sent, top10_Received, top10_Total)


# In[38]:


# Function to calculate the % of total delays per Section Block
def perc_sb(df_sb, totalSentMessageDelay, totalReceivedMessageDelay):
    # Sent
    for i in range(0,len(df_sb['NAME']),1):
        df_sb.iloc[i,13] = round((df_sb.iloc[i,4]/totalSentMessageDelay)*100,2)
        df_sb.iloc[i,13] = str(df_sb.iloc[i,13])+'%'

    # Received
    for i in range(0,len(df_sb['NAME']),1):
        df_sb.iloc[i,14] = round((df_sb.iloc[i,5]/totalReceivedMessageDelay)*100,2)
        df_sb.iloc[i,14] = str(df_sb.iloc[i,14])+'%'

    # Total
    for i in range(0,len(df_sb['NAME']),1):
        df_sb.iloc[i,15] = round((df_sb.iloc[i,6]/(totalSentMessageDelay+totalReceivedMessageDelay))*100,2)
        df_sb.iloc[i,15] = str(df_sb.iloc[i,15])+'%'
        
    return (df_sb)


# In[39]:


# Function to calculate the % of delays per locomotive
def perc_delay_loco(df_loco):
    df_loco.iloc[:,7] = (df_loco.iloc[:,6]/sum(df_loco.iloc[:,6]))*100
    df_loco.iloc[:,7] = round(df_loco.iloc[:,7],2)
    df_loco['%_DELAY'] = df_loco['%_DELAY'].astype(str)+'%'
    
    return(df_loco)


# In[40]:


# Function to retrieve month and year variables from user and verify if they are valid
def period():
    # Packages
    from datetime import date
    
    today = date.today()
    error = 0
    
    while(error == 0):      
        try:
            month = int(input("Enter month (1 or 2 digits): "))
            year = int(input("Enter year (4 digits): "))
        except: 
            month = 0    # Otherwise it will give error when going to 'if'
            year = 0
       
        
        if(not(month >= 1 and month <= 12)):
            print("\n!! ERROR - An incorrect value was entered. !!")
            print("Please try again.\n")
        elif(len(str(year))!=4):
            print("\n!! ERROR - An incorrect value was entered. !!")
            print("Please try again.\n")
        elif(year < 2020):
            print("\n!! ERROR - Please only enter the year 2020 or above. \n")
        elif(today < date(year, month, 1) ):
            print("\n!! ERROR - Please fill in past dates only. \n")
        else:
            error = 1    # Leaving the loop
            
    month = str(month)
    year = str(year)
    
    return(month, year)   


# In[41]:


# Function to convert train prefix to hexadecimal
def pref_hex(i, df_sent):
    # Transform the prefix from decimal to hexadecimal
    prefix = hex(int(df_sent.iloc[i,1]))
    prefix = prefix[2:10]
    prefix = prefix.upper()
    temp = ''              #empty string delcaration for temporary variable

    if len(prefix) == 2:
        prefix_hex = '-' + prefix + '-'
    elif len(prefix) == 1:
        prefix_hex = '-0' + prefix + '-'

    while len(prefix) > 2:     
        if len(prefix)%2 == 1:
            temp = temp + '-0' +  prefix[0] # take first character
            prefix = prefix[1:10] # take the rest
            prefix_hex =  temp + '-' + prefix + '-'
        elif len(prefix)%2 == 0:
            temp = temp + '-' + prefix[0:2] # take first two characters
            prefix = prefix[2:10] # take the rest
            prefix_hex = temp + '-' + prefix + '-'
    return(prefix_hex)


# In[42]:


def received_channel_analysis(df_received, mm_log, abr_log_received, month, year):
    import logging

    for i in range(0,len(df_received['HM_ID_HM']),1):    

            logging.info('MESSAGE RECEIVED:'+str(i))
            logging.info('HM_ID_HM:'+str(df_received.iloc[i,0]))
            (new_format, new_format_I, new_format_II) = mm_timeformat(i,df_received)
            (file_line) = mm_search_received(i, df_received, new_format, mm_log, new_format_I, new_format_II)
            remove_line(mm_log, file_line)
            (file_line) = mm_hex(file_line)
            (prefix_hex) = pref_hex(i, df_received)
            (gprs_attempt, inmarsat_attempt) = abr_channel_received(prefix_hex, file_line, abr_log_received)
            (df_received)= channel_count_received(i, df_received, gprs_attempt, inmarsat_attempt)
                # logging
            logging.info('FILE LINE: '+str(file_line))
            logging.info('PREFIX HEX: '+str(prefix_hex))
            logging.info('TCP: '+str(gprs_attempt)+' INMARSAT: '+str(inmarsat_attempt))
            logging.info('----------------------------------------------------------------')
            
    return(df_received)


# In[43]:


# Function to remove the line from the log that was used. This will avoid duplicated lines (with same macro, same loco and timestamp)
def remove_line(mm_log, file_line):
    mm_log.pop(mm_log.index(file_line)) 


# In[44]:


# Function to create Communication Report
def report(month, year, totalSentMessages, totalReceivedMessages, totalMessages, averageSentDelivery, averageReceivedDelivery, averageTotalDelivery, percentageSentDelay, totalSentDelay, percentageReceivedDelay, totalReceivedDelay, percentageTotalDelay, totalDelay, a, averageSentDeliveryDelay, averageReceivedDeliveryDelay, averageDeliveryDelay, gprs_delay_sent, inmarsat_delay_sent, gprs_queue_sent, inmarsat_queue_sent, gprs_non_queue_sent, inmarsat_non_queue_sent, no_attempt_sent, perc_gprs_delay_sent, perc_inmarsat_delay_sent, perc_no_attempt_sent, perc_gprs_queue_sent, perc_gprs_non_queue_sent, perc_inmarsat_queue_sent, perc_inmarsat_non_queue_sent, perc_total_queue_sent, perc_total_non_queue_sent, gprs_delay_received, inmarsat_delay_received, perc_gprs_delay_received, perc_inmarsat_delay_received, ts_total, df_loco, df_sb, images_location, report_location):
    import docx
    from docx2pdf import convert
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Creating the file
    mydoc = docx.Document()
    
    # Adding Title, month and year
    mydoc.add_heading("Communication Report", 0)
    mydoc.add_paragraph("Month: " + str(month) + " -  Year: " + str(year))
    
    # Adding section 'Overall Message Statistics'
    mydoc.add_heading("Overall Messages Statistics", 1)
    mydoc.add_paragraph("Total amount of Sent messages: " + str(totalSentMessages))
    mydoc.add_paragraph("Total amount of Received messages: " + str(totalReceivedMessages))
    mydoc.add_paragraph("Total amount of messages:  " + str(totalMessages))
    mydoc.add_paragraph("Average Delivery Time for Sent messages: " + str(averageSentDelivery))
    mydoc.add_paragraph("Average Delivery Time for Received messages: " + str(averageReceivedDelivery))
    mydoc.add_paragraph("Average Delivery Time of messages: " + str(averageTotalDelivery))
    mydoc.add_picture(images_location+"/density_"+str(month)+"_"+str(year)+".png")
    last_paragraph = mydoc.paragraphs[-1]     # Putting the image into the center
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adding section 'Delayed Messages Statistics'
    mydoc.add_heading("Delayed Messages Statistics ", 1)
    mydoc.add_paragraph("Time to consider a message to be delayed: " + str(a))
    mydoc.add_paragraph("Total amount of Sent messages w/ delay: " + str(totalSentDelay))
    mydoc.add_paragraph("Total amount of Received messages w/ delay: " + str(totalReceivedDelay))
    mydoc.add_paragraph("Total amount of messages w/ delay: " + str(totalDelay))
    mydoc.add_paragraph("Percentage of Sent messages w/ delay: " + str(percentageSentDelay))
    mydoc.add_paragraph("Percentage of Received messages w/ delay: " + str(percentageReceivedDelay))
    mydoc.add_paragraph("Percentage of messages w/ delay: " + str(percentageTotalDelay))
    mydoc.add_paragraph("Average Delivery Time for Sent messages w/ delay: " + str(averageSentDeliveryDelay))
    mydoc.add_paragraph("Average Delivery Time for Received messages w/ delay: " + str(averageReceivedDeliveryDelay))
    mydoc.add_paragraph("Average Delivery Time for messages w/ delay: " + str(averageDeliveryDelay))
    mydoc.add_paragraph("Number of Delays per Day: ")
    table = mydoc.add_table(rows=len(ts_total)+1, cols=4)
    cell = table.cell(0, 0)
    cell.text = 'DATE'
    cell = table.cell(0, 1)
    cell.text = 'SENT'
    cell = table.cell(0, 2)
    cell.text = 'RECEIVED'
    cell = table.cell(0, 3)
    cell.text = 'TOTAL'
    for i in range(1, len(ts_total)+1,1):
        for j in range(0,4,1):
            cell = table.cell(i, j)
            cell.text = str(ts_total.iloc[i-1,j])
    table.style = 'LightGrid-Accent1'
    mydoc.add_picture(images_location+"/time_series_"+str(month)+"_"+str(year)+".png")
    last_paragraph = mydoc.paragraphs[-1]     # Putting the image into the center
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER    
    
    # Adding section 'Channel Statistics'
    mydoc.add_heading("Channel Statistics", 1)
    mydoc.add_paragraph("Total Sent messages w/ delay by GPRS: " + str(gprs_delay_sent))
    mydoc.add_paragraph("Total Sent messages w/ delay by SATELLITE: " + str(inmarsat_delay_sent))
    mydoc.add_paragraph("Total Sent messages w/ delay with no attempt (timed out): " + str(no_attempt_sent))
    mydoc.add_paragraph("Percentage Sent message w/ delay by GPRS: " + str(perc_gprs_delay_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay by SATELLITE: " + str(perc_inmarsat_delay_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay with no attempt: " + str(perc_no_attempt_sent))
    mydoc.add_paragraph("Total Sent messages w/ delay by GPRS caused by QUEUE: " + str(gprs_queue_sent))
    mydoc.add_paragraph("Total Sent messages w/ delay by GPRS and actual delay (non-QUEUE): " + str(gprs_non_queue_sent))
    mydoc.add_paragraph("Total Sent messages w/ delay by SATELLITE caused by QUEUE: " + str(inmarsat_queue_sent))
    mydoc.add_paragraph("Total Sent messages w/ delay by SATELLITE and actual delay (non-QUEUE): " + str(inmarsat_non_queue_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay by GPRS caused by QUEUE: " + str(perc_gprs_queue_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay by GPRS not caused by QUEUE (actual delay): " + str(perc_gprs_non_queue_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay by SATELLITE caused by QUEUE: " + str(perc_inmarsat_queue_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay by SATELLITE not caused by QUEUE (actual delay): " + str(perc_inmarsat_non_queue_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay caused by QUEUE: " + str(perc_total_queue_sent))
    mydoc.add_paragraph("Percentage Sent messages w/ delay not caused by QUEUE (actual delay): " + str(perc_total_non_queue_sent))
    mydoc.add_paragraph("Total Received message w/ delay by GPRS: " + str(gprs_delay_received))
    mydoc.add_paragraph("Total Received messages w/ delay by SATELLITE: " + str(inmarsat_delay_received))
    mydoc.add_paragraph("Percentage Received message w/ delay by GPRS: " + str(perc_gprs_delay_received))
    mydoc.add_paragraph("Percentage Received messages w/ delay by SATELLITE: " + str(perc_inmarsat_delay_received))    
    
    # Adding section 'Locomotive Statistics'
    mydoc.add_heading("Locomotive Statistics", 1)
    mydoc.add_paragraph("All Locomotives ordered by 'Total Delays': ")
    df_loco = df_loco.sort_values(by=['D_TOTAL'], ascending = False)
    table = mydoc.add_table(rows=len(df_loco)+1, cols=7)
    cell = table.cell(0, 0)
    cell.text = 'LOCO ID'
    cell = table.cell(0, 1)
    cell.text = 'TOTAL MESSAGES SENT'
    cell = table.cell(0, 2)
    cell.text = 'TOTAL MESSAGES RECEIVED'
    cell = table.cell(0, 3)
    cell.text = 'TOTAL MESSAGES'
    cell = table.cell(0, 4)
    cell.text = 'DELAYS SENT'
    cell = table.cell(0, 5)
    cell.text = 'DELAYS RECEIVED'
    cell = table.cell(0, 6)
    cell.text = 'TOTAL DELAYS'
    for i in range(1, len(df_loco)+1,1):
        for j in range(0,7,1):
            cell = table.cell(i, j)
            cell.text = str(df_loco.iloc[i-1,j])
    table.style = 'LightGrid-Accent1'    
    mydoc.add_picture(images_location+"/bar_loco_"+str(month)+"_"+str(year)+".png")
    last_paragraph = mydoc.paragraphs[-1]     # Putting the image into the center
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Adding section 'Section Block Statistics'
    mydoc.add_heading("Section Block Statistics", 1)
    mydoc.add_paragraph("All Section Blocks ordered by 'Total Delays': ")
    df_sb.drop('SECTION_DEBUT', inplace=True, axis=1)
    df_sb.drop('CANTON', inplace=True, axis=1)
    df_sb.drop('SECTION_TETE', inplace=True, axis=1)
    df_sb = df_sb.sort_values(by=['COUNTS_TOTAL'], ascending = False)

    table = mydoc.add_table(rows=len(df_sb)+1, cols=4)
    cell = table.cell(0, 0)
    cell.text = 'SECTION BLOCK'
    cell = table.cell(0, 1)
    cell.text = 'DELAYS SENT MESSAGES'
    cell = table.cell(0, 2)
    cell.text = 'DELAYS RECEIVED MESSAGES'
    cell = table.cell(0, 3)
    cell.text = 'TOTAL DELAYS'
    for i in range(1, len(df_sb)+1,1):
        for j in range(0,4,1):
            cell = table.cell(i, j)
            cell.text = str(df_sb.iloc[i-1,j])
    table.style = 'LightGrid-Accent1'  
    mydoc.add_picture(images_location+"/bar_total_"+str(month)+"_"+str(year)+".png")
    last_paragraph = mydoc.paragraphs[-1]     # Putting the image into the center
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Saving the document in both .docx and .pdf
    mydoc.save(report_location + "/Communication Report "+str(month) + " " + str(year)+ ".docx")
    convert(report_location + "/Communication Report "+str(month) + " " + str(year) + ".docx")


# In[45]:


# Function to return total of delays per SB and total delivery time per SB for sent, received and total(sent + received) messages
def sb_calc(df_sb, df_sent, df_received, totalSentMessageDelay, totalReceivedMessageDelay):
    
    # Sent
    for i in range(0,len(df_sb['NAME']),1):
        for j in range(0,totalSentMessageDelay,1):
            if df_sent.iloc[j,3] == 1002 and df_sb.iloc[i,1] in df_sent.iloc[j,6]:
                df_sb.iloc[i,4] += 1
                df_sb.iloc[i,7] += df_sent.iloc[j,12]

            if df_sent.iloc[j,3] != 1002 and df_sb.iloc[i,2] in df_sent.iloc[j,6]:
                df_sb.iloc[i,4] +=1
                df_sb.iloc[i,7] += df_sent.iloc[j,12]

    # Received
    for i in range(0,len(df_sb['NAME']),1):
        for j in range(0,totalReceivedMessageDelay,1):
            if df_sb.iloc[i,1] in df_received.iloc[j,6]:
                df_sb.iloc[i,5] += 1
                df_sb.iloc[i,8] += df_received.iloc[j,12]

            if df_sb.iloc[i,2] in df_received.iloc[j,6]:
                df_sb.iloc[i,5] += 1
                df_sb.iloc[i,8] += df_received.iloc[j,12]

            if df_sb.iloc[i,3] in df_received.iloc[j,6]:
                df_sb.iloc[i,5] += 1
                df_sb.iloc[i,8] += df_received.iloc[j,12]

    # !! There may be a slight difference between totalReceivedMessageDelay and the sum of df_sb['COUNTS'].
    # This is because the script does not consider messages on switches 
    # There are also some messages that have 'Canton:' but don't contain an actual location !!

    # Total
    for i in range(0, len(df_sb['NAME']),1):
        df_sb.iloc[i,6] = df_sb.iloc[i,4] + df_sb.iloc[i,5]
        df_sb.iloc[i,9] = df_sb.iloc[i,7] + df_sb.iloc[i,8]

    # Fix for Received Messages Count
    x=0
    for i in range(0,len(df_sb['NAME']),1):
        x = x + df_sb.iloc[i,5]

    # A new value is returned to 'totalReceivedMessageDelay' to correctly count the % column.
    # Please refer to the difference in the comments of the previous section. (Between '!!...!!')
    totalReceivedMessageDelay = int(x)
    
    return (df_sb,totalReceivedMessageDelay)


# In[46]:


# Main function for channel analysis
def sent_channel_analysis(df_sent, mm_log, abr_log_sent, month, year):
    import logging

    for i in range(0,len(df_sent['HM_ID_HM']),1):    
        try:
            logging.info('MESSAGE SENT:'+str(i))
            logging.info('MT_ID_MT:'+str(df_sent.iloc[i,11]))
            (new_format, new_format_I, new_format_II) = mm_timeformat(i,df_sent)
            (file_line) = mm_search_sent(i, df_sent, new_format, mm_log, new_format_I, new_format_II)
            remove_line(mm_log, file_line)
            (file_line) = mm_hex(file_line)
            (prefix_hex) = pref_hex(i, df_sent)
            (gprs_attempt, inmarsat_attempt) = abr_channel_sent(prefix_hex, file_line, abr_log_sent)
            (df_sent)= channel_count_sent(i, df_sent, gprs_attempt, inmarsat_attempt)
            # logging
            logging.info('FILE LINE: '+str(file_line))
            logging.info('PREFIX HEX: '+str(prefix_hex))
            logging.info('TCP: '+str(gprs_attempt)+' INMARSAT: '+str(inmarsat_attempt))
            logging.info('----------------------------------------------------------------')

        except Exception as e:
            logging.ERROR(e)
              # If an error occurs, the channel values are being set as NULL.
            df_sent.iloc[i,17] += 1 


    return(df_sent)


# In[47]:


# Function for converting time series data into a dataframe
def time_series_conv(ts_sent, ts_received):
    # Packages
    import pandas as pd
    
    # Converting Pandas Series to Dataframe
    ts_sent = ts_sent.to_frame('SENT')
    ts_received = ts_received.to_frame('RECEIVED')
    # Creating and formatting new dataframe
    ts_total = pd.merge(ts_sent,ts_received,on='TS_DATE')
    ts_total['TOTAL'] = 0
    ts_total.iloc[:,2] = ts_total.iloc[:,0] + ts_total.iloc[:,1]
    ts_total.index.name = 'TS_DATE'
    ts_total.reset_index(inplace=True)
    
    return(ts_total)


# In[48]:


# Function to decompose the time series
def ts_decompose(ts_sent, ts_received, month, year, images_location):
    # Packages
    from statsmodels.tsa.seasonal import seasonal_decompose
    import matplotlib.pyplot as plt
    
    # Creating the decomposition variable
    decomp_sent = seasonal_decompose(ts_sent, period=12) # Periods can be thought of as special cases of intervals
    decomp_received = seasonal_decompose(ts_received, period=12)

    # Trend
    tend_sent = decomp_sent.trend
    tend_received = decomp_received.trend
    # Seasonal
    seas_sent = decomp_sent.seasonal
    seas_received = decomp_received.seasonal
    # Random
    rand_sent = decomp_sent.resid
    rand_received = decomp_received.resid

    plt.figure(1)
    plt.subplot(2,2,1)
    plt.plot(ts_sent,color ='cornflowerblue', label='Sent')
    plt.plot(ts_received,color ='tab:red',label='Received')
    plt.title('Time Series')
    plt.xticks([])
    plt.yticks([])
    plt.subplot(2,2,2)
    plt.plot(tend_sent,color ='cornflowerblue')
    plt.plot(tend_received,color ='tab:red')
    plt.xticks([])
    plt.yticks([])
    plt.title('Trend')
    plt.subplot(2,2,3)
    plt.plot(seas_sent,color ='cornflowerblue')
    plt.plot(seas_received,color ='tab:red')
    plt.xticks([])
    plt.yticks([])
    plt.title('Seasonality')
    #plt.subplot(2,2,4)
    #plt.plot(rand_sent,color ='cornflowerblue')
    #plt.plot(rand_received,color ='tab:red')
    #plt.xticks([])
    #plt.yticks([])
    #plt.title('Resid')
    plt.savefig(images_location+'/time_series_decompose'+'_'+ month + '_' + year +'.png')


# In[49]:


# Function to create time series graph for sent and receives messages
def ts_graph(ts_sent, ts_received, month, year, images_location):
    # Packages 
    import matplotlib.pyplot as plt
    
    # Sent
    plt.plot(ts_sent,color ='cornflowerblue', label='Sent')
    plt.xticks(rotation=90)
    plt.title("Delays per Day")

    # Received
    plt.plot(ts_received,color ='tab:red', label='Received')
    plt.xticks(rotation=90)
    plt.title("Delays per Day")
    
    for i in range(0,max(ts_sent),100):
        plt.axhline(y=i, color='black', linestyle='--',linewidth=0.5)
    plt.tight_layout()
    plt.legend(loc='best')
    plt.savefig(images_location+'/time_series'+'_'+ month + '_' + year +'.png')


# In[50]:


# Function to prepare for time series
def ts_prep(df_sent, df_received):
    # Packages
    import pandas as pd
    
    #Sent
    df_sent['TS_DATE'] = 0
    df_sent.iloc[:,13] = df_sent.iloc[:,5].astype(str) # adding a new column and convert to string (pandas .astype)
    df_sent.iloc[:,13] = df_sent.iloc[:,13].str.slice(start=0,stop=10)
    # Grouping number of messages per day
    ts_sent = df_sent.groupby(['TS_DATE']).size()
   
    #Received
    df_received['TS_DATE'] = 0
    df_received.iloc[:,13] = df_received.iloc[:,5].astype(str) # adding a new column and convert to string (pandas .astype)
    df_received.iloc[:,13] = df_received.iloc[:,13].str.slice(start=0,stop=10)
    # Grouping number of messages per day
    ts_received = df_received.groupby(['TS_DATE']).size()
    
    return(df_sent, df_received, ts_sent, ts_received)


# In[51]:


# Function for variables
def variables_decl(month, year):
    #Packages
    from datetime import timedelta
    from datetime import datetime
    import logging
    # Variable declaration
    a = timedelta(seconds=300)

    
    return (a)


# In[52]:


# Function to prepare images for database insertion
def var_prep(month, year):

    density = 'density_'+month+'_'+year+'.png'
    dispersion_all = 'dispersion_all_'+month+'_'+year+'.png'
    dispersion_sent_received = 'dispersion_sent_received_'+month+'_'+year+'.png'
    dispersion_total = 'dispersion_total_'+month+'_'+year+'.png'
    bar_all = 'bar_all_'+month+'_'+year+'.png'
    bar_sent = 'bar_sent_'+month+'_'+year+'.png'
    bar_received = 'bar_received_'+month+'_'+year+'.png'
    bar_total = 'bar_total_'+month+'_'+year+'.png'
    bar_loco = 'bar_loco_'+month+'_'+year+'.png'
    time_series = 'time_series_'+month+'_'+year+'.png' 
    time_series_decompose = 'time_series_decompose_'+month+'_'+year+'.png' 

    month = int(month)
    year = int(year)
    
    return(density, dispersion_all, dispersion_sent_received, dispersion_total, bar_all, bar_sent, bar_received, bar_total, bar_loco, time_series, time_series_decompose, month, year)

